#!/usr/bin/env python3

import os
import json
import logging
import time
import sys
import pandas as pd
from typing import Dict, Any, List, Optional

# Import configuration
try:
    from config import CONFIG
except ImportError:
    # Default configuration if config.py is not available
    CONFIG = {
        "output_dir": "./agent_outputs",
        "memory_dir": "./agent_memory",
        "default_model": "openrouter/quasar-alpha",
        "api_key": "your-api-key",
        "endpoint": "https://openrouter.ai/api/v1/chat/completions",
        "memory_db": "agent_memory.db",
        "sqlite_db": "test_sqlite.db"
    }
    # Ensure output directories exist
    os.makedirs(CONFIG["output_dir"], exist_ok=True)
    os.makedirs(CONFIG["memory_dir"], exist_ok=True)

# Import agent framework, assuming it's in the same directory or in PYTHONPATH
try:
    from agent_framework import AgentSystem, Agent, DynamicAgent
    from agent_tools import ToolRegistry
except ImportError:
    print("Agent framework modules not found. Make sure agent_framework.py and agent_tools.py are in your PYTHONPATH.")
    sys.exit(1)

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(CONFIG["output_dir"], "sap_analysis.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Define custom tools for SAP data analysis
def read_csv(filename: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Read a CSV file and return its contents.
    
    Args:
        filename: The name of the CSV file to read
        context: Additional context for reading the file
        
    Returns:
        Dict with the CSV data and metadata
    """
    logger.info(f"Reading CSV file: {filename}")
    
    try:
        # Read the CSV file
        df = pd.read_csv(filename)
        
        # Basic statistics and metadata
        row_count = len(df)
        column_count = len(df.columns)
        columns = list(df.columns)
        column_types = {col: str(df[col].dtype) for col in df.columns}
        
        # Get sample data (first 5 rows)
        sample_data = df.head(5).to_dict(orient='records')
        
        # Get basic statistics
        numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
        statistics = {}
        if numeric_columns:
            statistics = {
                col: {
                    'min': float(df[col].min()) if not pd.isna(df[col].min()) else None,
                    'max': float(df[col].max()) if not pd.isna(df[col].max()) else None,
                    'mean': float(df[col].mean()) if not pd.isna(df[col].mean()) else None,
                    'median': float(df[col].median()) if not pd.isna(df[col].median()) else None
                }
                for col in numeric_columns
            }
        
        # Missing values information
        missing_values = df.isnull().sum().to_dict()
        
        return {
            "filename": filename,
            "row_count": row_count,
            "column_count": column_count,
            "columns": columns,
            "column_types": column_types,
            "sample_data": sample_data,
            "statistics": statistics,
            "missing_values": missing_values,
            "dataframe": df  # Include the dataframe for further processing
        }
    
    except Exception as e:
        logger.error(f"Error reading CSV file {filename}: {str(e)}")
        return {
            "filename": filename,
            "error": str(e),
            "status": "error"
        }

def read_docx(filename: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Read a DOCX file and extract its content.
    
    Args:
        filename: The name of the DOCX file to read
        context: Additional context for reading the file
        
    Returns:
        Dict with the DOCX content
    """
    logger.info(f"Reading DOCX file: {filename}")
    
    try:
        import docx
        
        # Read the DOCX file
        doc = docx.Document(filename)
        
        # Extract text content
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "filename": filename,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
            "status": "success"
        }
    
    except ImportError:
        logger.error("python-docx package not installed. Install with: pip install python-docx")
        return {
            "filename": filename,
            "error": "python-docx package not installed",
            "status": "error"
        }
    except Exception as e:
        logger.error(f"Error reading DOCX file {filename}: {str(e)}")
        return {
            "filename": filename,
            "error": str(e),
            "status": "error"
        }

def analyze_data_quality(data: Dict[str, Any], context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Analyze the quality of data in a dataframe.
    
    Args:
        data: The data to analyze (must contain a 'dataframe' key)
        context: Additional context for the analysis
        
    Returns:
        Dict with data quality metrics
    """
    logger.info("Analyzing data quality")
    
    try:
        df = data.get('dataframe')
        if df is None:
            return {
                "error": "No dataframe provided",
                "status": "error"
            }
        
        # Calculate completeness (percentage of non-missing values)
        total_cells = df.shape[0] * df.shape[1]
        missing_cells = df.isnull().sum().sum()
        completeness = (total_cells - missing_cells) / total_cells if total_cells > 0 else 0
        
        # Check for duplicates
        duplicate_rows = df.duplicated().sum()
        
        # Check data types consistency
        type_consistency = {}
        for col in df.columns:
            if df[col].dtype == 'object':
                # For string columns, check if they contain mixed data types
                try:
                    pd.to_numeric(df[col], errors='raise')
                    type_consistency[col] = "Mixed (numbers stored as strings)"
                except:
                    # Check if it contains dates
                    try:
                        pd.to_datetime(df[col], errors='raise')
                        type_consistency[col] = "Mixed (dates stored as strings)"
                    except:
                        type_consistency[col] = "Consistent (string)"
            else:
                type_consistency[col] = "Consistent (numeric)"
        
        # Detect outliers in numeric columns
        outliers = {}
        numeric_cols = df.select_dtypes(include=['number']).columns
        for col in numeric_cols:
            # Use IQR method to detect outliers
            q1 = df[col].quantile(0.25)
            q3 = df[col].quantile(0.75)
            iqr = q3 - q1
            lower_bound = q1 - (1.5 * iqr)
            upper_bound = q3 + (1.5 * iqr)
            outliers[col] = ((df[col] < lower_bound) | (df[col] > upper_bound)).sum()
        
        # Overall data quality score (simple average of components)
        # Completeness: 0-1 score already calculated
        # Accuracy proxy (outliers): percentage of non-outlier values
        total_numeric_values = sum(len(df[col]) for col in numeric_cols)
        total_outliers = sum(outliers.values())
        accuracy = 1 - (total_outliers / total_numeric_values if total_numeric_values > 0 else 0)
        
        # Consistency proxy: percentage of columns with consistent data types
        consistent_cols = sum(1 for status in type_consistency.values() if status.startswith("Consistent"))
        consistency = consistent_cols / len(type_consistency) if type_consistency else 0
        
        # Timeliness: We don't have date of data collection, assume 0.8 for demo
        timeliness = 0.8
        
        # Return quality metrics
        return {
            "quality_metrics": {
                "completeness": round(completeness, 2),
                "accuracy": round(accuracy, 2),
                "consistency": round(consistency, 2),
                "timeliness": round(timeliness, 2),
                "overall": round((completeness + accuracy + consistency + timeliness) / 4, 2)
            },
            "missing_values": data.get('missing_values', {}),
            "duplicate_rows": int(duplicate_rows),
            "type_consistency": type_consistency,
            "outliers": outliers,
            "status": "success"
        }
    
    except Exception as e:
        logger.error(f"Error analyzing data quality: {str(e)}")
        return {
            "error": str(e),
            "status": "error"
        }

def list_cognitive_patterns(prompt: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    List available cognitive planning patterns.
    
    Args:
        prompt: The prompt requesting pattern information
        context: Additional context
        
    Returns:
        Dict with available cognitive patterns
    """
    logger.info("Listing cognitive planning patterns")
    
    # Simulated cognitive patterns
    patterns = {
        "tot": {
            "id": "tot",
            "name": "Tree of Thoughts",
            "description": "Explores multiple solution paths in a structured tree format, evaluating branches and selecting optimal paths.",
            "best_for": ["Complex decision-making", "Multiple viable alternatives", "Business optimization"],
            "stages": ["Problem framing", "Branch generation", "Branch evaluation", "Path selection", "Solution development"]
        },
        "mad": {
            "id": "mad",
            "name": "Multi-Agent Debate",
            "description": "Simulates a debate between agents with different perspectives to thoroughly explore a problem.",
            "best_for": ["Controversial topics", "Multi-stakeholder problems", "Nuanced policy decisions"],
            "stages": ["Topic framing", "Position statements", "Arguments and rebuttals", "Criticism", "Synthesis"]
        },
        "mcr": {
            "id": "mcr",
            "name": "Metacognitive Reflection",
            "description": "Analyzes and improves the problem-solving process itself, detecting biases and exploring alternatives.",
            "best_for": ["Bias detection", "Decision quality improvement", "Learning from past decisions"],
            "stages": ["Initial solution", "Process analysis", "Bias detection", "Alternative exploration", "Process optimization"]
        },
        "col": {
            "id": "col",
            "name": "Collaborative Chains",
            "description": "Coordinates multiple specialized agents in a workflow to solve complex problems.",
            "best_for": ["Multi-domain problems", "Specialized expertise requirements", "Process workflows"],
            "stages": ["Task decomposition", "Agent assignment", "Parallel processing", "Result integration", "Quality verification"]
        }
    }
    
    # Based on the context, recommend a pattern
    recommended_pattern = "tot"  # Default recommendation
    reasoning = "Tree of Thoughts is recommended for business optimization problems as it allows structured exploration of multiple solution paths with rigorous evaluation."
    
    return {
        "status": "success",
        "patterns": patterns,
        "recommendation": recommended_pattern,
        "reasoning": reasoning
    }

def create_cognitive_session(prompt: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Create a cognitive planning session.
    
    Args:
        prompt: The prompt requesting session creation
        context: Additional context
        
    Returns:
        Dict with session information
    """
    logger.info("Creating cognitive planning session")
    
    # Extract the pattern ID from the context or default to "tot"
    pattern_id = "tot"
    if context and "cognitive_pattern_explorer" in context:
        pattern_data = context["cognitive_pattern_explorer"].get("output", {})
        if isinstance(pattern_data, dict) and "recommendation" in pattern_data:
            pattern_id = pattern_data["recommendation"]
    
    # Generate a session ID
    import uuid
    session_id = str(uuid.uuid4())
    
    # Get pattern information
    patterns = list_cognitive_patterns("", {})["patterns"]
    pattern_info = patterns.get(pattern_id, {})
    
    # Extract problem description from context
    problem_description = "Optimize SAP business processes based on data analysis"
    if context and "business_challenge_framer" in context:
        challenge_data = context["business_challenge_framer"].get("output", {})
        if isinstance(challenge_data, dict) and "key_challenges" in challenge_data:
            challenges = challenge_data["key_challenges"]
            if challenges:
                problem_description = "; ".join(challenges[:3])  # Use first 3 challenges
    
    # Create session info
    return {
        "status": "success",
        "session_id": session_id,
        "pattern": pattern_info.get("name", "Tree of Thoughts"),
        "pattern_id": pattern_id,
        "description": problem_description,
        "stages": pattern_info.get("stages", []),
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    }

def get_next_cognitive_step(session_id: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Get the next step in a cognitive planning session.
    
    Args:
        session_id: The ID of the cognitive session
        context: Additional context
        
    Returns:
        Dict with next step information
    """
    logger.info(f"Getting next cognitive step for session {session_id}")
    
    # Check context for completed steps
    completed_steps = []
    current_step = None
    
    if context and "tot_workflow_navigator" in context:
        workflow_data = context["tot_workflow_navigator"].get("output", {})
        if isinstance(workflow_data, dict) and "completed_steps" in workflow_data:
            completed_steps = workflow_data["completed_steps"]
    
    # Get pattern stages from the session creation
    stages = []
    if context and "cognitive_session_creator" in context:
        session_data = context["cognitive_session_creator"].get("output", {})
        if isinstance(session_data, dict) and "stages" in session_data:
            stages = session_data["stages"]
    
    # Determine current step
    if not stages:
        # Default stages for Tree of Thoughts
        stages = ["Problem framing", "Branch generation", "Branch evaluation", "Path selection", "Solution development"]
    
    # Find the first uncompleted stage
    for stage in stages:
        if stage not in completed_steps:
            current_step = stage
            break
    
    # If all steps are completed, return completed status
    if not current_step:
        return {
            "status": "completed",
            "message": "All cognitive planning steps have been completed",
            "session_id": session_id
        }
    
    # Define instructions for the current step
    step_instructions = {
        "Problem framing": "Define the key aspects of the business problem based on SAP data analysis. Identify dimensions to optimize.",
        "Branch generation": "Generate 3-5 distinct approaches to address the identified business problems.",
        "Branch evaluation": "Evaluate each approach based on feasibility, impact, cost, and alignment with business objectives.",
        "Path selection": "Select the most promising approach based on the evaluation.",
        "Solution development": "Develop a detailed solution plan for the selected approach."
    }
    
    return {
        "status": "in_progress",
        "session_id": session_id,
        "current_step": current_step,
        "instructions": step_instructions.get(current_step, "Proceed with the next step in the cognitive process."),
        "completed_steps": completed_steps,
        "remaining_steps": [s for s in stages if s not in completed_steps and s != current_step]
    }

def select_cognitive_action(action_id: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Select an action in a cognitive planning session.
    
    Args:
        action_id: The ID of the action to select
        context: Additional context
        
    Returns:
        Dict with action selection result
    """
    logger.info(f"Selecting cognitive action: {action_id}")
    
    # Extract session ID from context
    session_id = ""
    if context and "cognitive_session_creator" in context:
        session_data = context["cognitive_session_creator"].get("output", {})
        if isinstance(session_data, dict) and "session_id" in session_data:
            session_id = session_data["session_id"]
    
    # Extract current step from context
    current_step = ""
    if context and "tot_workflow_navigator" in context:
        workflow_data = context["tot_workflow_navigator"].get("output", {})
        if isinstance(workflow_data, dict) and "current_step" in workflow_data:
            current_step = workflow_data["current_step"]
    
    return {
        "status": "success",
        "session_id": session_id,
        "step": current_step,
        "selected_action": action_id,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }

def submit_cognitive_result(result: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Submit a result for a step in a cognitive planning session.
    
    Args:
        result: The result to submit
        context: Additional context
        
    Returns:
        Dict with submission result
    """
    logger.info("Submitting cognitive result")
    
    # Extract session ID and current step from context
    session_id = ""
    current_step = ""
    completed_steps = []
    
    if context and "cognitive_session_creator" in context:
        session_data = context["cognitive_session_creator"].get("output", {})
        if isinstance(session_data, dict) and "session_id" in session_data:
            session_id = session_data["session_id"]
    
    if context and "tot_workflow_navigator" in context:
        workflow_data = context["tot_workflow_navigator"].get("output", {})
        if isinstance(workflow_data, dict):
            if "current_step" in workflow_data:
                current_step = workflow_data["current_step"]
            if "completed_steps" in workflow_data:
                completed_steps = workflow_data["completed_steps"]
    
    # Add current step to completed steps if not already present
    if current_step and current_step not in completed_steps:
        completed_steps.append(current_step)
    
    return {
        "status": "success",
        "session_id": session_id,
        "step_completed": current_step,
        "completed_steps": completed_steps,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }

def get_cognitive_results(session_id: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Get the results of a cognitive planning session.
    
    Args:
        session_id: The ID of the cognitive session
        context: Additional context
        
    Returns:
        Dict with session results
    """
    logger.info(f"Getting cognitive results for session {session_id}")
    
    # In a real implementation, this would retrieve results from storage
    # For this demo, we'll generate simulated results based on context
    
    # Default results
    results = {
        "problem_statement": "Optimize SAP business processes based on data analysis",
        "selected_approach": "Comprehensive process optimization",
        "implementation_steps": [
            "Data quality improvement",
            "Process efficiency analysis",
            "System configuration updates",
            "User training and adoption",
            "Performance monitoring"
        ],
        "expected_outcomes": [
            "Improved data accuracy and completeness",
            "Streamlined business processes",
            "Reduced operational costs",
            "Enhanced reporting capabilities",
            "Better decision-making support"
        ]
    }
    
    # Try to extract more specific information from context
    if context:
        # Extract business challenges
        if "business_challenge_framer" in context:
            challenge_data = context["business_challenge_framer"].get("output", {})
            if isinstance(challenge_data, dict):
                if "key_challenges" in challenge_data:
                    results["challenges"] = challenge_data["key_challenges"]
                if "opportunities" in challenge_data:
                    results["opportunities"] = challenge_data["opportunities"]
        
        # Extract cost-benefit analysis
        if "cost_benefit_analyzer" in context:
            cost_benefit_data = context["cost_benefit_analyzer"].get("output", {})
            if isinstance(cost_benefit_data, dict) and "improvement_areas" in cost_benefit_data:
                results["improvement_areas"] = cost_benefit_data["improvement_areas"]
    
    return {
        "status": "success",
        "session_id": session_id,
        "pattern": "Tree of Thoughts",
        "results": results,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }

def transform_cognitive_results(format_type: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Transform cognitive results into a specified format.
    
    Args:
        format_type: The type of transformation to apply
        context: Additional context
        
    Returns:
        Dict with transformed results
    """
    logger.info(f"Transforming cognitive results to format: {format_type}")
    
    # Extract session results from context
    results = {}
    session_id = ""
    
    if context and "cognitive_results_transformer" in context:
        transform_data = context["cognitive_results_transformer"].get("tools", {})
        if isinstance(transform_data, dict) and "cognitive:get_results" in transform_data:
            results = transform_data["cognitive:get_results"].get("results", {})
            session_id = transform_data["cognitive:get_results"].get("session_id", "")
    
    # Transform based on format type
    if format_type == "research":
        # Format for research presentation
        transformed = {
            "research_title": "SAP Business Process Optimization Research",
            "methodology": "Tree of Thoughts Cognitive Analysis",
            "key_findings": [
                "Data quality issues impact operational efficiency",
                "Process optimization opportunities exist in order processing",
                "Integration improvements needed between systems",
                "Training and adoption are critical success factors"
            ],
            "recommendations": results.get("implementation_steps", []),
            "expected_impact": results.get("expected_outcomes", [])
        }
    elif format_type == "executive":
        # Format for executive presentation
        transformed = {
            "executive_summary": "SAP Business Process Optimization Strategy",
            "business_impact": "High",
            "investment_required": "Medium",
            "roi_timeline": "12-18 months",
            "strategic_alignment": "Strong",
            "key_recommendations": results.get("implementation_steps", [])
        }
    else:
        # Default format
        transformed = results
    
    return {
        "status": "success",
        "original_format": "cognitive_results",
        "target_format": format_type,
        "session_id": session_id,
        "transformed_results": transformed,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }

# Register tools
ToolRegistry.register_tool("file:read_csv", read_csv)
ToolRegistry.register_tool("file:read_docx", read_docx)
ToolRegistry.register_tool("data:analyze_quality", analyze_data_quality)
ToolRegistry.register_tool("cognitive:list_patterns", list_cognitive_patterns)
ToolRegistry.register_tool("cognitive:create_session", create_cognitive_session)
ToolRegistry.register_tool("cognitive:get_next_step", get_next_cognitive_step)
ToolRegistry.register_tool("cognitive:select_action", select_cognitive_action)
ToolRegistry.register_tool("cognitive:submit_result", submit_cognitive_result)
ToolRegistry.register_tool("cognitive:get_results", get_cognitive_results)
ToolRegistry.register_tool("cognitive:transform_for_research", transform_cognitive_results)

class SAPDataWorkflow(AgentSystem):
    """
    SAP Data Analysis Workflow System.
    
    This system implements the workflow defined in the JSON configuration
    to analyze SAP data, identify optimization opportunities, and create
    implementation plans.
    """
    
    def __init__(self):
        """Initialize the SAP Data Workflow System."""
        super().__init__("SAP Data Workflow System")
        
        # Add all agents from the workflow
        self._configure_agents()
    
    def _configure_agents(self):
        """Configure all agents for the workflow."""
        # Data inventory agent
        self.add_agent(Agent(
            agent_id="data_inventory_agent",
            content="Create an inventory of what we can expect to find in the sap_data.csv file based on the column names provided. Identify the different sections (SalesOrders, SalesOrderItems, BusinessPartners, etc.) and key fields in each section.",
            tools=["file:read_csv"],
            output_format={
                "type": "json",
                "schema": {
                    "data_sections": ["string"],
                    "key_fields": ["string"],
                    "analysis_approach": "string"
                }
            }
        ))
        
        # Data quality validator
        self.add_agent(Agent(
            agent_id="data_quality_validator",
            content="Perform a comprehensive data quality assessment on the sap_data.csv file. Analyze completeness, accuracy, consistency, and timeliness of the data. Identify missing values, duplicates, outliers, and any other data quality issues that could impact analysis.",
            tools=["file:read_csv", "data:analyze_quality", "planning:tree_of_thought"],
            read_from=["data_inventory_agent"],
            output_format={
                "type": "json",
                "schema": {
                    "quality_metrics": {
                        "completeness": "number",
                        "accuracy": "number",
                        "consistency": "number",
                        "timeliness": "number"
                    },
                    "data_issues": ["string"],
                    "critical_gaps": ["string"],
                    "recommendations": ["string"]
                }
            }
        ))
        
        # Data integrity checker
        self.add_agent(Agent(
            agent_id="data_integrity_checker",
            content="Analyze referential integrity across different sections of the sap_data.csv file. Verify that foreign key relationships (like SALESORDERID between orders and items, or PARTNERID between orders and business partners) are maintained. Identify any integrity violations or data inconsistencies.",
            tools=["file:read_csv", "planning:tree_of_thought"],
            read_from=["data_inventory_agent", "data_quality_validator"],
            output_format={
                "type": "json",
                "schema": {
                    "verified_relationships": ["string"],
                    "integrity_violations": ["string"],
                    "data_consistency_issues": ["string"],
                    "impact_on_analysis": "string"
                }
            }
        ))
        
        # Sales orders analyzer
        self.add_agent(Agent(
            agent_id="sales_orders_analyzer",
            content="Analyze the sales orders data in the sap_data.csv file. Focus on order headers including fields like SALESORDERID, CREATEDBY, PARTNERID, GROSSAMOUNT, etc. Provide a summary of key metrics and patterns. Consider the data quality issues identified previously in your analysis.",
            tools=["file:read_csv"],
            read_from=["data_inventory_agent", "data_quality_validator", "data_integrity_checker"],
            output_format={
                "type": "json",
                "schema": {
                    "order_count": "number",
                    "key_metrics": "object",
                    "insights": "string",
                    "quality_adjusted_analysis": "string"
                }
            }
        ))
        
        # Sales order items analyzer
        self.add_agent(Agent(
            agent_id="sales_order_items_analyzer",
            content="Analyze the sales order items data in the sap_data.csv file. Focus on line items including fields like SALESORDERID, SALESORDERITEM, PRODUCTID, QUANTITY, etc. Identify relationships with order headers. Adjust your analysis based on the previously identified data quality issues.",
            tools=["file:read_csv"],
            read_from=["data_inventory_agent", "sales_orders_analyzer", "data_quality_validator", "data_integrity_checker"],
            output_format={
                "type": "json",
                "schema": {
                    "item_count": "number",
                    "key_metrics": "object",
                    "insights": "string",
                    "quality_adjusted_analysis": "string"
                }
            }
        ))
        
        # Business partners analyzer
        self.add_agent(Agent(
            agent_id="business_partners_analyzer",
            content="Analyze the business partners data in the sap_data.csv file. Focus on partner information including fields like PARTNERID, PARTNERROLE, COMPANYNAME, etc. Identify relationships with sales orders. Consider data quality limitations in your analysis.",
            tools=["file:read_csv"],
            read_from=["data_inventory_agent", "sales_orders_analyzer", "data_quality_validator", "data_integrity_checker"],
            output_format={
                "type": "json",
                "schema": {
                    "partner_count": "number",
                    "key_metrics": "object",
                    "insights": "string",
                    "quality_adjusted_analysis": "string"
                }
            }
        ))
        
        # Supply chain analyzer
        self.add_agent(Agent(
            agent_id="supply_chain_analyzer",
            content="Analyze the supply chain data in the sap_data.csv file. Focus on relevant fields like SupplierID, ProductID, CustomerID, ShipDate, etc. Identify connections to sales data. Adjust analysis based on known data quality issues.",
            tools=["file:read_csv"],
            read_from=["data_inventory_agent", "sales_orders_analyzer", "sales_order_items_analyzer", "data_quality_validator", "data_integrity_checker"],
            output_format={
                "type": "json",
                "schema": {
                    "record_count": "number",
                    "key_metrics": "object",
                    "insights": "string",
                    "quality_adjusted_analysis": "string"
                }
            }
        ))
        
        # Industry benchmark researcher
        self.add_agent(Agent(
            agent_id="industry_benchmark_researcher",
            content="Research industry benchmarks and standards for key SAP business metrics including order processing efficiency, supply chain performance, customer relationship management, and sales performance. Contextualizing our findings against industry standards will provide valuable perspective.",
            tools=["planning:tree_of_thought"],
            output_format={
                "type": "json",
                "schema": {
                    "sales_benchmarks": "object",
                    "order_processing_benchmarks": "object",
                    "supply_chain_benchmarks": "object",
                    "partner_management_benchmarks": "object",
                    "data_quality_benchmarks": "object",
                    "benchmark_sources": ["string"]
                }
            }
        ))
        
        # Benchmarked analysis integrator
        self.add_agent(Agent(
            agent_id="benchmarked_analysis_integrator",
            content="Integrate the industry benchmarks with our SAP data analysis to contextualize findings. Compare our metrics against industry standards and identify areas where performance is above or below expectations. This contextual analysis will enhance the value of our insights.",
            tools=["planning:tree_of_thought"],
            read_from=["industry_benchmark_researcher", "sales_orders_analyzer", "sales_order_items_analyzer", "business_partners_analyzer", "supply_chain_analyzer"],
            output_format={
                "type": "json",
                "schema": {
                    "contextualized_metrics": "object",
                    "performance_gaps": ["string"],
                    "competitive_advantages": ["string"],
                    "industry_position": "string"
                }
            }
        ))
        
        # Reports document analyzer
        self.add_agent(Agent(
            agent_id="reports_document_analyzer",
            content="Analyze the Sales Reports.docx document to identify relevant SAP reports for our data. Match report types to the data we've analyzed so far and determine which standard SAP reports would best address the identified performance gaps.",
            tools=["file:read_docx"],
            read_from=["sales_orders_analyzer", "sales_order_items_analyzer", "business_partners_analyzer", "supply_chain_analyzer", "benchmarked_analysis_integrator"],
            output_format={
                "type": "json",
                "schema": {
                    "available_reports": ["string"],
                    "recommended_reports": ["string"],
                    "gap_addressing_reports": ["string"],
                    "insights": "string"
                }
            }
        ))
        
        # Business challenge framer
        self.add_agent(Agent(
            agent_id="business_challenge_framer",
            content="Based on all the SAP data analyzed so far, frame the key business challenges and opportunities that need to be addressed. This will serve as input to our cognitive planning process. Incorporate both data quality limitations and industry benchmark comparisons in your assessment.",
            read_from=["sales_orders_analyzer", "sales_order_items_analyzer", "business_partners_analyzer", "supply_chain_analyzer", "reports_document_analyzer", "data_quality_validator", "benchmarked_analysis_integrator"],
            output_format={
                "type": "json",
                "schema": {
                    "key_challenges": ["string"],
                    "opportunities": ["string"],
                    "priority_areas": ["string"],
                    "data_limitations": ["string"]
                }
            }
        ))
        
        # Cognitive pattern explorer
        self.add_agent(Agent(
            agent_id="cognitive_pattern_explorer",
            content="List all available cognitive planning patterns to identify the optimal approach for analyzing our SAP data. Focus on patterns that allow for structured exploration of multiple solution paths, especially those suited for business optimization in the context of data quality challenges.",
            tools=["cognitive:list_patterns"],
            output_format={
                "type": "json",
                "schema": {
                    "status": "string",
                    "patterns": "object",
                    "recommendation": "string",
                    "reasoning": "string"
                }
            }
        ))
        
        # Cognitive session creator
        self.add_agent(Agent(
            agent_id="cognitive_session_creator",
            content="Create a cognitive planning session using the Tree of Thoughts pattern to address the business challenges identified from our SAP data analysis. Use pattern_id parameter explicitly. Ensure the problem description includes awareness of data quality limitations and industry benchmarking context.",
            tools=["cognitive:create_session"],
            read_from=["cognitive_pattern_explorer", "business_challenge_framer"],
            output_format={
                "type": "json",
                "schema": {
                    "status": "string",
                    "session_id": "string",
                    "pattern": "string",
                    "description": "string",
                    "stages": ["string"]
                }
            }
        ))
        
        # Tree of Thoughts workflow navigator
        self.add_agent(Agent(
            agent_id="tot_workflow_navigator",
            content="Navigate through the Tree of Thoughts cognitive planning workflow to identify optimal business improvement approaches for the analyzed SAP data. Use the get_next_step, select_action, and submit_result functions to progress through the process. Ensure that data quality limitations are factored into the decision-making process.",
            tools=["cognitive:get_next_step", "cognitive:select_action", "cognitive:submit_result"],
            read_from=["cognitive_session_creator", "data_quality_validator"],
            output_format={
                "type": "json",
                "schema": {
                    "session_id": "string",
                    "completed_steps": ["string"],
                    "key_decisions": ["string"],
                    "final_status": "string"
                }
            }
        ))
        
        # Cognitive results transformer
        self.add_agent(Agent(
            agent_id="cognitive_results_transformer",
            content="Transform the completed Tree of Thoughts cognitive planning session into a structured business optimization plan for the SAP data. Extract specific recommendations and implementation approaches. Factor in data quality considerations and industry benchmarks to ensure recommendations are practical and impactful.",
            tools=["cognitive:get_results", "cognitive:transform_for_research"],
            read_from=["tot_workflow_navigator", "data_quality_validator", "benchmarked_analysis_integrator"],
            output_format={
                "type": "json",
                "schema": {
                    "business_focus": "string",
                    "optimization_plan": {
                        "goal": "string",
                        "key_approaches": ["string"],
                        "implementation_steps": ["string"]
                    },
                    "expected_benefits": ["string"],
                    "data_quality_dependencies": ["string"],
                    "cognitive_pattern_used": "string"
                }
            }
        ))
        
        # Cost benefit analyzer
        self.add_agent(Agent(
            agent_id="cost_benefit_analyzer",
            content="Conduct a comprehensive cost-benefit analysis for each of the potential improvement areas identified in the cognitive planning process. Estimate implementation costs (technology, process change, training), projected benefits (revenue increase, cost reduction, efficiency gains), and calculate ROI and payback periods to prioritize recommendations.",
            tools=["planning:chain_of_thought"],
            read_from=["cognitive_results_transformer", "benchmarked_analysis_integrator"],
            output_format={
                "type": "json",
                "schema": {
                    "improvement_areas": [
                        {
                            "area": "string",
                            "estimated_costs": "object",
                            "projected_benefits": "object",
                            "roi": "number",
                            "payback_period": "string",
                            "risk_factors": ["string"]
                        }
                    ],
                    "prioritized_recommendations": ["string"],
                    "financial_impact_summary": "string"
                }
            }
        ))
        
        # SAP implementation specialist
        self.add_agent(Agent(
            agent_id="sap_implementation_specialist",
            content="Provide detailed SAP-specific implementation guidance for the prioritized recommendations. Include SAP module configurations, transaction codes, tables, BAPIs, and other SAP-specific technical elements required for implementation. Ensure recommendations align with SAP best practices and standard SAP functionality.",
            tools=["planning:chain_of_thought"],
            read_from=["cognitive_results_transformer", "cost_benefit_analyzer", "reports_document_analyzer"],
            output_format={
                "type": "json",
                "schema": {
                    "implementation_plans": [
                        {
                            "recommendation": "string",
                            "sap_modules": ["string"],
                            "transaction_codes": ["string"],
                            "configuration_steps": ["string"],
                            "technical_requirements": ["string"],
                            "testing_approaches": ["string"]
                        }
                    ],
                    "implementation_timeline": "object",
                    "resource_requirements": "object",
                    "critical_success_factors": ["string"]
                }
            }
        ))
        
        # Data quality remediation planner
        self.add_agent(Agent(
            agent_id="data_quality_remediation_planner",
            content="Develop a detailed plan to address the data quality issues identified during the analysis. Focus on improving data completeness, accuracy, consistency, and timeliness to support successful implementation of the business improvement recommendations.",
            tools=["planning:chain_of_thought"],
            read_from=["data_quality_validator", "data_integrity_checker", "cognitive_results_transformer", "sap_implementation_specialist"],
            output_format={
                "type": "json",
                "schema": {
                    "data_quality_objectives": ["string"],
                    "remediation_steps": ["string"],
                    "data_governance_recommendations": ["string"],
                    "validation_procedures": ["string"],
                    "impact_on_business_improvements": "string"
                }
            }
        ))
        
        # Implementation risk analyzer
        self.add_agent(Agent(
            agent_id="implementation_risk_analyzer",
            content="Analyze potential risks and challenges for implementing the recommended SAP improvements, especially considering the identified data quality issues. Provide mitigation strategies to ensure successful implementation.",
            tools=["planning:chain_of_thought"],
            read_from=["cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "data_quality_remediation_planner"],
            output_format={
                "type": "json",
                "schema": {
                    "implementation_risks": [
                        {
                            "risk": "string",
                            "probability": "number",
                            "impact": "number",
                            "mitigation_strategy": "string"
                        }
                    ],
                    "critical_dependencies": ["string"],
                    "risk_adjusted_timeline": "object",
                    "contingency_recommendations": ["string"]
                }
            }
        ))
        
        # Dynamic implementation selector
        self.add_agent(DynamicAgent(
            agent_id="dynamic_implementation_selector",
            content="Select the most urgent implementation focus area based on all analysis",
            initial_prompt="Based on all SAP data analysis, industry benchmarking, cost-benefit analysis, and the Tree of Thoughts optimization planning, determine which business area requires most urgent implementation focus. Choose from: 'sales_optimization', 'customer_management', 'supply_chain_efficiency', or 'order_processing_improvement'. Consider both business impact and implementation feasibility in your selection.",
            read_from=["sales_orders_analyzer", "sales_order_items_analyzer", "business_partners_analyzer", "supply_chain_analyzer", "cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "implementation_risk_analyzer"],
            output_format={
                "type": "json",
                "schema": {
                    "selected_approach": "string",
                    "justification": "string",
                    "expected_impact": "string",
                    "implementation_complexity": "string"
                }
            },
            actions={
                "sales_optimization": {
                    "agent": "sales_optimization_planner",
                    "content": "Develop a detailed implementation plan for sales optimization based on the Tree of Thoughts analysis and SAP data insights. Include specific actionable recommendations with timelines, expected outcomes, SAP-specific implementation steps, and data quality dependencies.",
                    "readFrom": ["sales_orders_analyzer", "sales_order_items_analyzer", "cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "data_quality_remediation_planner"]
                },
                "customer_management": {
                    "agent": "customer_management_planner",
                    "content": "Develop a detailed implementation plan for customer relationship management based on the Tree of Thoughts analysis and SAP data insights. Include specific actionable recommendations with timelines, expected outcomes, SAP-specific implementation steps, and data quality dependencies.",
                    "readFrom": ["business_partners_analyzer", "sales_orders_analyzer", "cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "data_quality_remediation_planner"]
                },
                "supply_chain_efficiency": {
                    "agent": "supply_chain_planner",
                    "content": "Develop a detailed implementation plan for supply chain efficiency improvements based on the Tree of Thoughts analysis and SAP data insights. Include specific actionable recommendations with timelines, expected outcomes, SAP-specific implementation steps, and data quality dependencies.",
                    "readFrom": ["supply_chain_analyzer", "sales_order_items_analyzer", "cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "data_quality_remediation_planner"]
                },
                "order_processing_improvement": {
                    "agent": "order_processing_planner",
                    "content": "Develop a detailed implementation plan for order processing improvements based on the Tree of Thoughts analysis and SAP data insights. Include specific actionable recommendations with timelines, expected outcomes, SAP-specific implementation steps, and data quality dependencies.",
                    "readFrom": ["sales_orders_analyzer", "sales_order_items_analyzer", "cognitive_results_transformer", "cost_benefit_analyzer", "sap_implementation_specialist", "data_quality_remediation_planner"]
                }
            }
        ))
        
        # Implementation reliability tester
        self.add_agent(Agent(
            agent_id="implementation_reliability_tester",
            content="Design test scenarios and validation approaches to ensure reliable execution of the implementation plan. Identify potential execution bottlenecks and develop failsafe mechanisms to maintain implementation continuity.",
            tools=["planning:chain_of_thought"],
            read_from=["dynamic_implementation_selector", "sap_implementation_specialist", "implementation_risk_analyzer"],
            output_format={
                "type": "json",
                "schema": {
                    "test_scenarios": ["string"],
                    "validation_criteria": ["string"],
                    "performance_metrics": ["string"],
                    "reliability_recommendations": ["string"],
                    "contingency_procedures": ["string"]
                }
            }
        ))
        
        # Implementation validator
        self.add_agent(Agent(
            agent_id="implementation_validator",
            content="Validate the selected implementation plan against the original SAP data analysis, Tree of Thoughts cognitive results, industry benchmarks, and cost-benefit analysis. Identify any gaps, conflicts, or enhancement opportunities. Confirm that data quality dependencies are properly addressed.",
            read_from=["dynamic_implementation_selector", "cognitive_results_transformer", "sales_orders_analyzer", "sales_order_items_analyzer", "business_partners_analyzer", "supply_chain_analyzer", "benchmarked_analysis_integrator", "cost_benefit_analyzer", "data_quality_remediation_planner"],
            output_format={
                "type": "json",
                "schema": {
                    "validation_status": "string",
                    "alignment_score": "number",
                    "gaps_identified": ["string"],
                    "enhancement_opportunities": ["string"],
                    "data_dependency_validation": "string"
                }
            }
        ))
        
        # Comprehensive report creator
        self.add_agent(Agent(
            agent_id="comprehensive_report_creator",
            content="Create a comprehensive report that integrates all findings from the SAP data analysis, data quality assessment, industry benchmarking, Tree of Thoughts cognitive planning, cost-benefit analysis, and implementation planning. Provide a clear roadmap for business optimization with SAP-specific implementation details.",
            read_from=["*"],
            output_format={
                "type": "markdown",
                "sections": [
                    "Executive Summary",
                    "SAP Data Analysis Overview",
                    "Data Quality Assessment",
                    "Industry Benchmark Comparison",
                    "Business Challenges & Opportunities",
                    "Tree of Thoughts Cognitive Analysis",
                    "Strategic Optimization Plan",
                    "Cost-Benefit Analysis",
                    "SAP-Specific Implementation Details",
                    "Implementation Roadmap",
                    "Data Quality Remediation Plan",
                    "Risk Management",
                    "Expected Business Outcomes",
                    "Key Performance Indicators",
                    "Next Steps",
                    "Appendix: Detailed SAP Reports"
                ]
            }
        ))
        
        logger.info(f"Configured {len(self.agents)} agents for the SAP Data Workflow System")
    
    def run_analysis(self, sap_data_file: str, sales_reports_file: str = None) -> Dict[str, Any]:
        """
        Run the complete SAP data analysis workflow.
        
        Args:
            sap_data_file: Path to the SAP data CSV file
            sales_reports_file: Path to the Sales Reports DOCX file (optional)
            
        Returns:
            Dict with analysis results and report
        """
        logger.info(f"Starting SAP data analysis workflow for file: {sap_data_file}")
        
        # Prepare inputs
        inputs = {
            "sap_data.csv": sap_data_file
        }
        
        if sales_reports_file:
            inputs["Sales Reports.docx"] = sales_reports_file
        
        # Execute the workflow starting with data inventory
        results = self.execute_workflow("data_inventory_agent", inputs)
        
        # Extract the comprehensive report
        report = {}
        if "comprehensive_report_creator" in results and results["comprehensive_report_creator"]["status"] == "success":
            report = results["comprehensive_report_creator"]["output"]
        
        # Determine the selected implementation approach
        selected_approach = "unknown"
        if "dynamic_implementation_selector" in results and results["dynamic_implementation_selector"]["status"] == "success":
            selected_approach = results["dynamic_implementation_selector"].get("selected_action", "unknown")
        
        # Prepare the summary of results
        summary = {
            "analysis_complete": True,
            "agents_executed": list(results.keys()),
            "data_quality": {},
            "selected_implementation_approach": selected_approach,
            "report": report
        }
        
        # Add data quality information if available
        if "data_quality_validator" in results and results["data_quality_validator"]["status"] == "success":
            summary["data_quality"] = results["data_quality_validator"]["output"].get("quality_metrics", {})
        
        # Save the full results to a file
        results_filename = os.path.join(CONFIG["output_dir"], "sap_analysis_results.json")
        with open(results_filename, 'w') as f:
            json.dump(results, f, indent=2)
        
        logger.info(f"SAP data analysis workflow completed. Results saved to: {results_filename}")
        
        # Save the report to a file
        if report:
            report_filename = os.path.join(CONFIG["output_dir"], "sap_analysis_report.md")
            if isinstance(report, dict) and "content" in report:
                with open(report_filename, 'w') as f:
                    f.write(report["content"])
            elif isinstance(report, str):
                with open(report_filename, 'w') as f:
                    f.write(report)
            
            logger.info(f"SAP analysis report saved to: {report_filename}")
        
        return {
            "summary": summary,
            "results_file": results_filename,
            "report_file": report_filename if report else None
        }

def print_section(title: str):
    """Print a formatted section title."""
    print("\n" + "=" * 80)
    print(f" {title} ".center(80, "="))
    print("=" * 80)

def main():
    print_section("SAP Data Analysis Workflow")
    
    # Check if SAP data file is provided as argument
    import argparse
    parser = argparse.ArgumentParser(description='Run SAP data analysis workflow')
    parser.add_argument('--sap-data', default='sap_data.csv', help='Path to SAP data CSV file')
    parser.add_argument('--sales-reports', default=None, help='Path to Sales Reports DOCX file (optional)')
    args = parser.parse_args()
    
    # Check if files exist
    sap_data_exists = os.path.isfile(args.sap_data)
    sales_reports_exists = args.sales_reports and os.path.isfile(args.sales_reports)
    
    if not sap_data_exists:
        print(f"Warning: SAP data file '{args.sap_data}' not found.")
        print("Creating sample SAP data file for demonstration...")
        create_sample_sap_data(args.sap_data)
    
    if args.sales_reports and not sales_reports_exists:
        print(f"Warning: Sales Reports file '{args.sales_reports}' not found.")
        print("Creating sample Sales Reports file for demonstration...")
        create_sample_sales_reports(args.sales_reports)
    
    # Create and run the SAP data workflow
    workflow = SAPDataWorkflow()
    
    print(f"Processing SAP data file: {args.sap_data}")
    if args.sales_reports:
        print(f"Using Sales Reports file: {args.sales_reports}")
    
    # Run the analysis
    results = workflow.run_analysis(args.sap_data, args.sales_reports)
    
    # Display results summary
    print_section("Analysis Results Summary")
    
    print(f"Analysis completed: {results['summary']['analysis_complete']}")
    print(f"Selected implementation approach: {results['summary']['selected_implementation_approach']}")
    
    if results['summary']['data_quality']:
        print("\nData Quality Metrics:")
        for metric, value in results['summary']['data_quality'].items():
            print(f"  {metric.capitalize()}: {value}")
    
    print(f"\nResults file: {results['results_file']}")
    if results['report_file']:
        print(f"Report file: {results['report_file']}")
    
    print("\nNote: This is a simulated analysis. In a real implementation, the agents would ")
    print("perform actual data analysis and generate detailed reports.")

def create_sample_sap_data(filename: str):
    """Create a sample SAP data CSV file for demonstration purposes."""
    import csv
    
    # Define sample data
    headers = [
        "RECORDTYPE", "SALESORDERID", "SALESORDERITEM", "PARTNERID", "PRODUCTID", 
        "QUANTITY", "QUANTITYUNIT", "GROSSAMOUNT", "CURRENCY", "CREATEDBY", 
        "CREATEDAT", "COMPANYNAME", "PARTNERROLE", "DELIVERYDATE", "SUPPLIERNAME"
    ]
    
    sample_data = [
        ["HEADER", "1001", "", "2001", "", "", "", "1200.00", "USD", "ADMIN", "2023-01-15", "ABC Company", "CUSTOMER", "", ""],
        ["ITEM", "1001", "0010", "", "3001", "2", "EA", "400.00", "USD", "", "", "", "", "2023-01-20", "Supplier A"],
        ["ITEM", "1001", "0020", "", "3002", "4", "EA", "800.00", "USD", "", "", "", "", "2023-01-20", "Supplier B"],
        ["HEADER", "1002", "", "2002", "", "", "", "950.00", "USD", "ADMIN", "2023-01-16", "XYZ Corp", "CUSTOMER", "", ""],
        ["ITEM", "1002", "0010", "", "3003", "1", "EA", "950.00", "USD", "", "", "", "", "2023-01-22", "Supplier C"],
        ["HEADER", "1003", "", "2001", "", "", "", "500.00", "USD", "USER1", "2023-01-17", "ABC Company", "CUSTOMER", "", ""],
        ["ITEM", "1003", "0010", "", "3001", "1", "EA", "200.00", "USD", "", "", "", "", "2023-01-25", "Supplier A"],
        ["ITEM", "1003", "0020", "", "3004", "1", "EA", "300.00", "USD", "", "", "", "", "2023-01-25", "Supplier D"],
        ["PARTNER", "", "", "2001", "", "", "", "", "", "", "2023-01-01", "ABC Company", "CUSTOMER", "", ""],
        ["PARTNER", "", "", "2002", "", "", "", "", "", "", "2023-01-01", "XYZ Corp", "CUSTOMER", "", ""],
        ["PARTNER", "", "", "4001", "", "", "", "", "", "", "2023-01-01", "Supplier A", "SUPPLIER", "", ""],
        ["PARTNER", "", "", "4002", "", "", "", "", "", "", "2023-01-01", "Supplier B", "SUPPLIER", "", ""],
        ["PARTNER", "", "", "4003", "", "", "", "", "", "", "2023-01-01", "Supplier C", "SUPPLIER", "", ""],
        ["PARTNER", "", "", "4004", "", "", "", "", "", "", "2023-01-01", "Supplier D", "SUPPLIER", "", ""],
        ["PRODUCT", "", "", "", "3001", "", "EA", "200.00", "USD", "", "2023-01-01", "", "", "", "Supplier A"],
        ["PRODUCT", "", "", "", "3002", "", "EA", "200.00", "USD", "", "2023-01-01", "", "", "", "Supplier B"],
        ["PRODUCT", "", "", "", "3003", "", "EA", "950.00", "USD", "", "2023-01-01", "", "", "", "Supplier C"],
        ["PRODUCT", "", "", "", "3004", "", "EA", "300.00", "USD", "", "2023-01-01", "", "", "", "Supplier D"]
    ]
    
    # Write to CSV file
    try:
        with open(filename, 'w', newline='') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(headers)
            writer.writerows(sample_data)
        
        print(f"Sample SAP data created at: {filename}")
    except Exception as e:
        print(f"Error creating sample data: {str(e)}")

def create_sample_sales_reports(filename: str):
    """Create a sample Sales Reports DOCX file for demonstration purposes."""
    try:
        # Try to import python-docx
        import docx
        from docx.shared import Pt
        
        # Create a new document
        doc = docx.Document()
        
        # Add a title
        doc.add_heading('SAP Sales Reports Documentation', 0)
        
        # Add introduction
        doc.add_paragraph('This document provides information about standard SAP sales reports and their usage for business analysis.')
        
        # Add sales report section
        doc.add_heading('Standard SAP Sales Reports', 1)
        
        reports = [
            ("S_ALR_87012993", "Sales Order Analysis", "Analyzes sales orders by customer, product, and time period."),
            ("S_ALR_87012939", "Sales Revenue Analysis", "Provides revenue breakdown by various dimensions."),
            ("S_ALR_87013070", "Customer Analysis", "Analyzes customer performance and relationship metrics."),
            ("S_ALR_87013532", "Product Analysis", "Provides insights into product performance and sales."),
            ("S_ALR_87013611", "Supply Chain Performance", "Analyzes delivery performance and supply chain metrics.")
        ]
        
        # Add report details
        for txcode, name, desc in reports:
            p = doc.add_paragraph()
            p.add_run(f"{name} ({txcode})").bold = True
            doc.add_paragraph(desc)
        
        # Add a table with sample KPIs
        doc.add_heading('Key Performance Indicators', 1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'KPI Name'
        header_cells[1].text = 'Description'
        header_cells[2].text = 'Related Report'
        
        # Add KPIs
        kpis = [
            ("Order Fulfillment Rate", "Percentage of orders fulfilled completely and on time", "S_ALR_87012993"),
            ("Average Order Value", "Average monetary value of sales orders", "S_ALR_87012939"),
            ("Customer Retention Rate", "Percentage of customers making repeat purchases", "S_ALR_87013070"),
            ("Product Return Rate", "Percentage of products returned by customers", "S_ALR_87013532"),
            ("Delivery Time", "Average time from order to delivery", "S_ALR_87013611")
        ]
        
        for kpi, desc, report in kpis:
            row_cells = table.add_row().cells
            row_cells[0].text = kpi
            row_cells[1].text = desc
            row_cells[2].text = report
        
        # Save the document
        doc.save(filename)
        print(f"Sample Sales Reports document created at: {filename}")
        
    except ImportError:
        # If python-docx is not available, create a text file instead
        with open(filename, 'w') as f:
            f.write("SAP Sales Reports Documentation\n\n")
            f.write("Standard SAP Sales Reports:\n")
            f.write("- S_ALR_87012993: Sales Order Analysis\n")
            f.write("- S_ALR_87012939: Sales Revenue Analysis\n")
            f.write("- S_ALR_87013070: Customer Analysis\n")
            f.write("- S_ALR_87013532: Product Analysis\n")
            f.write("- S_ALR_87013611: Supply Chain Performance\n")
        
        print(f"Sample Sales Reports text file created at: {filename} (python-docx not available)")
    except Exception as e:
        print(f"Error creating sample Sales Reports: {str(e)}")

if __name__ == "__main__":
    main()
